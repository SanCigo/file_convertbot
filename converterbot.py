import json
import logging
import os
import time
from pprint import pprint
import requests
import telegram 
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, MessageHandler, \
    CallbackQueryHandler, PicklePersistence, Filters, ConversationHandler
import re
from docx import Document
from loutils import doc2pdf


if not telegram.__version__.startswith("13."):
    print("This bot only runs on 13.x version of the library. 13.15 reccomended.")
    exit()


script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)


# put the token in the same path inside a text file
with open("token.txt", "r") as f:
    TOKEN = f.read()

# users with access to restricted features
with open("premiums.txt", "r") as f:
    PREMIUMS = [int(i) for i in f.readlines()]  # list of telegram ids separated by newline


def build_menu(buttons,
               n_cols,
               header_buttons=None,
               footer_buttons=None):
    menu = [buttons[i:i + n_cols] for i in range(0, len(buttons), n_cols)]
    if header_buttons:
        menu.insert(0, [header_buttons])
    if footer_buttons:
        menu.append([footer_buttons])
    return menu


my_persistence = PicklePersistence("converter.pickle")
updater = Updater(token=TOKEN, use_context=True, persistence=my_persistence)
dispatcher = updater.dispatcher
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
ASK_TARGET, ASK_FILE, CONVERT = range(3)
ASK_FILENAME, ASK_TITLE, ASK_CONTENT = range(3)


def api_id(update, context):
    context.user_data["api_id"] = context.args[0]
    update.message.reply_markdown("*ID settato correttamente.*\nUtilizza /start per iniziare a convertire i file.")


def ask_for_category(update, context):
    if "api_id" in context.user_data:
        context.user_data["category_target"] = []
        categories = [
            InlineKeyboardButton("🗄️ Archivio", callback_data="1archive"),
            InlineKeyboardButton("🎧 Audio", callback_data="1audio"),
            InlineKeyboardButton("📐 AutoCAD", callback_data="1cad"),
            InlineKeyboardButton("📄 Documento", callback_data="1document"),
            InlineKeyboardButton("📕 eBook", callback_data="1ebook"),
            InlineKeyboardButton("#️⃣ Hash", callback_data="1hash"),
            InlineKeyboardButton("🖼️Immagine", callback_data="1image"),
            InlineKeyboardButton("📊 Metadati", callback_data="1metadata"),
            InlineKeyboardButton("🎬 Video", callback_data="1video")
        ]
        reply_markup = InlineKeyboardMarkup(build_menu(categories, n_cols=2))
        if update.callback_query:
            context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                          message_id=update.callback_query.message.message_id,
                                          text="Scegli una categoria.",
                                          reply_markup=reply_markup)
        else:
            context.bot.send_message(chat_id=update.message.chat_id, text="Scegli una categoria.",
                                     reply_markup=reply_markup)
        return ASK_TARGET
    else:
        update.message.reply_markdown("*Benvenuto in File Converter Bot!*\nPer iniziare, hai bisogno di una *chiave "
                                      "API*. Scopri come ottenerla a [questo link]("
                                      "https://www.api2convert.com/documentation/getting-started#getting-an-api-key).\nRicorda, "
                                      "*non condividere la tua chiave API* con nessuno ad eccezione "
                                      "del bot. Una volta ottenuta, impostala con ``` /set chiave_api ```")
        return ConversationHandler.END


def ask_for_target(update, context):
    with open("conversions_parsed.json") as f:
        dict_conversioni = json.load(f)
    targets_buttons = []
    if len(context.user_data["category_target"]) > 1:
        targets = dict_conversioni[context.user_data["category_target"][0]]
        print(targets)
        context.user_data["category_target"].pop(1)
        context.user_data["flag_done"] = False
    else:
        targets = dict_conversioni[update.callback_query.data[1:]]
        print(targets)
        context.user_data["category_target"].append(update.callback_query.data[1:])
    for t in targets:
        targets_buttons.append(InlineKeyboardButton(t, callback_data=t))
    print(targets_buttons)
    reply_markup = InlineKeyboardMarkup(build_menu(targets_buttons, n_cols=3,
                                                   footer_buttons=InlineKeyboardButton("🔙 Indietro",
                                                                                       callback_data="2")))
    context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                  message_id=update.callback_query.message.message_id, text="Scegli un target.",
                                  reply_markup=reply_markup)
    return ASK_FILE


def ask_for_file(update, context):
    context.user_data["category_target"].append(update.callback_query.data)
    context.user_data["flag_done"] = True
    back_button = [InlineKeyboardButton("🔙 Indietro", callback_data="1")]
    reply_markup = InlineKeyboardMarkup(build_menu(back_button, n_cols=1))
    context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                  message_id=update.callback_query.message.message_id, text="Ora invia il file da "
                                                                                            "convertire. *Dimensioni "
                                                                                            "massime: 20 MB*.",
                                  reply_markup=reply_markup, parse_mode=telegram.ParseMode.MARKDOWN)
    return CONVERT


def convert(update, context):
    if "flag_done" in context.user_data and context.user_data["flag_done"]:
        if update.message.photo and update.message.photo[-1].file_size < 20000000 and not update.message.video:
            file = update.message.photo[-1].get_file()
            nome_file = file.file_path.rsplit("/", 1)[1]
            file.download(nome_file)
            file = open(nome_file, "rb")
        elif update.message.effective_attachment.file_size < 20000000:
            file = update.message.effective_attachment.get_file()
            nome_file = update.message.effective_attachment.file_name
            file.download(nome_file)
            file = open(nome_file, "rb")
        else:
            update.message.reply_markdown("Il file supera il limite di *20 MB*.")
            return ConversationHandler.END
        file_dict = {"file": file}
        url_get_job = "https://api2.online-convert.com/jobs"
        api_key = {'x-oc-api-key': context.user_data["api_id"]}
        category_and_target = "{\"conversion\": [{\"category\": \"" + context.user_data["category_target"][0] + \
                              "\", \"target\": \"" + context.user_data["category_target"][1] + "\"}], \"process\": " \
                                                                                               "true}"
        response_job = requests.post(url_get_job, headers=api_key, data=category_and_target)
        # pprint(response_job.json())
        temp_dict_upload = response_job.json()
        job_id = temp_dict_upload["id"]
        context.user_data[job_id] = nome_file
        url_upload = temp_dict_upload["server"] + "/upload-file/" + job_id
        api_key['x-oc-upload-uuid'] = str(id(file))
        requests.post(url_upload, headers=api_key, files=file_dict)
        file.close()
        os.remove(nome_file)
        update.message.reply_text("File caricato con successo. Inizio conversione... (più il file è grande, più tempo "
                                  "sarà richiesto)")
        url_get_job += "/" + job_id
        while True:
            response_job_status = requests.get(url_get_job, headers=api_key)
            json_status = response_job_status.json()
            # pprint(json_status)
            if json_status["status"]["code"] == "completed":
                response_file = requests.get(json_status["output"][0]["uri"])
                update.message.reply_text("Completato.")

                context.bot.send_chat_action(chat_id=update.effective_chat.id,
                                             action=telegram.ChatAction.UPLOAD_DOCUMENT)
                final_name = right_filename(nome_file, json_status["output"][0]["filename"])
                with open(final_name, "wb") as f:
                    f.write(response_file.content)
                with open(final_name, "rb") as f:
                    context.bot.send_document(chat_id=update.effective_chat.id, document=f)
                os.remove(final_name)
                break
            time.sleep(3)
        return ConversationHandler.END
    else:
        update.message.reply_text("Seleziona prima una categoria ed un target. Inizia con /start")


def right_filename(nome_originario, nuovo_nome):
    nome = nome_originario.rsplit(".", 1)
    estensione = nuovo_nome.rsplit(".", 1)
    print(nome, estensione)
    return nome[0] + "." + estensione[1]


def privacy(update, context):
    update.message.reply_markdown("[Norma sulla privacy](https://telegra.ph/Normativa-sulla-privacy-03-20)")


"""From here on, premiums user features only"""

def ask_pdfdocx(update, context):
    if update.effective_chat.id in PREMIUMS:
        cancel_button = [InlineKeyboardButton("🔵 Word", callback_data="4Word"),
                         InlineKeyboardButton("🔴 PDF", callback_data="4PDF"),
                         InlineKeyboardButton("❌ Annulla", callback_data="3")]
        reply_markup = InlineKeyboardMarkup(build_menu(cancel_button, n_cols=2))
        if not update.callback_query:
            message = context.bot.send_message(chat_id=update.effective_chat.id, text="Seleziona il *tipo di file.*",
                                               reply_markup=reply_markup, parse_mode=telegram.ParseMode.MARKDOWN)
            context.user_data["message_id"] = message.message_id
        else:
            context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                          message_id=update.callback_query.message.message_id,
                                          text="Seleziona il *tipo di file.*",
                                          reply_markup=reply_markup,
                                          parse_mode=telegram.ParseMode.MARKDOWN)
        return ASK_FILENAME


def ask_filename(update, context):
    context.user_data["message_id"] = update.callback_query.message.message_id
    context.user_data["type"] = update.callback_query.data
    back_button = [InlineKeyboardButton("🔙 Indietro", callback_data="0")]
    reply_markup = InlineKeyboardMarkup(build_menu(back_button, n_cols=1))
    update.effective_message.edit_text("Scrivi ora il *nome del file."
                                       "\nFormato:* " + update.callback_query.data[1:],
                                       reply_markup=reply_markup, parse_mode=telegram.ParseMode.MARKDOWN)
    return ASK_TITLE


def ask_title(update, context):
    context.user_data["contents"] = []
    context.user_data["first_time"] = True
    if not update.callback_query:
        is_name_ok = re.search(r"([a-zA-Z0-9\s_\\.\-\(\):])$", update.message.text)
        if is_name_ok:
            context.user_data["filename"] = update.message.text
        else:
            update.message.reply_markdown("*Il nome del file non è valido.* Non deve contenere *; / ,*")
            return ASK_TITLE
    message_id = context.user_data["message_id"]
    back_button = [InlineKeyboardButton("🔙 Indietro", callback_data=context.user_data["type"])]
    reply_markup = InlineKeyboardMarkup(build_menu(back_button, n_cols=1))
    context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                  text="Scrivi ora il *titolo del documento."
                                       "\nFormato:* " + context.user_data["type"][1:] +
                                       "\n*Nome del file:* " + context.user_data["filename"],
                                  message_id=message_id,
                                  reply_markup=reply_markup,
                                  parse_mode=telegram.ParseMode.MARKDOWN)
    return ASK_CONTENT


def ask_content(update, context):
    if not update.callback_query:
        back_button = [InlineKeyboardButton("🔙 Indietro", callback_data="6")]
        if not context.user_data["first_time"]:
            if update.message.photo:
                file = update.message.photo[-1].get_file()
                nome_file = file.file_path.rsplit("/", 1)[1]
                file.download(nome_file)
                context.user_data["contents"].append("1MG3" + nome_file)
            else:
                context.user_data["contents"].append(update.message.text)
            text = "Aggiunto con successo. Al momento hai inserito *" + str(
                len(context.user_data["contents"])) + "* paragrafi. Quando hai finito, usa /done "
        else:
            text = "Titolo aggiunto con successo. Al momento hai inserito *0* paragrafi. Quando hai finito, usa /done"
            context.user_data["title"] = update.message.text
        context.user_data["first_time"] = False
    else:
        if len(context.user_data["contents"]) == 1:
            print("HEY")
            back_button = [InlineKeyboardButton("🔙 Indietro", callback_data="5")]
            context.user_data["contents"].pop()
        else:
            back_button = [InlineKeyboardButton("🔙 Indietro", callback_data="6")]
            context.user_data["contents"].pop()
        text = "Al momento hai inserito *{0}* paragrafi. Quando hai finito, usa /done ".format(str(
            len(context.user_data["contents"])))
        reply_markup = InlineKeyboardMarkup(build_menu(back_button, n_cols=1))
        message = context.bot.edit_message_text(chat_id=update.effective_chat.id,
                                                message_id=context.user_data["message_id"],
                                                text=text,
                                                reply_markup=reply_markup,
                                                parse_mode=telegram.ParseMode.MARKDOWN)
        context.user_data["message_id"] = message.message_id
        return ASK_CONTENT
    reply_markup = InlineKeyboardMarkup(build_menu(back_button, n_cols=1))
    message = context.bot.send_message(chat_id=update.effective_chat.id,
                                       text=text,
                                       reply_markup=reply_markup,
                                       parse_mode=telegram.ParseMode.MARKDOWN)
    context.user_data["message_id"] = message.message_id
    return ASK_CONTENT


def indietro(update, context):
    data = update.callback_query.data
    if data == "2":
        ask_pdfdocx(update, context)
    elif data[0] == "4":
        ask_filename(update, context)
        return ASK_TITLE
    elif data == "5":
        ask_title(update, context)
        return ASK_CONTENT
    elif data == "6":
        ask_content(update, context)
        return ASK_CONTENT


def done(update, context):
    update.message.reply_markdown("Fine. Costruisco il file...")
    titolo = context.user_data["title"]
    nome_file = context.user_data["filename"] + ".docx"
    formato = context.user_data["type"]
    paragrafi = context.user_data["contents"]
    document = Document()
    document.add_heading(titolo, 0)
    for p in paragrafi:
        if p[0:4] == "1MG3":
            document.add_picture(p[4:])
        else:
            document.add_paragraph(p)
    path_nome_file = "Files/" + nome_file
    document.save(path_nome_file)
    if formato == "4PDF":
        doc2pdf(path_nome_file)
        os.remove(path_nome_file)
        nome_file = nome_file.replace(".docx", ".pdf")
        path_nome_file = "Files/" + nome_file
    with open(path_nome_file, "rb") as f:
        context.bot.send_chat_action(chat_id=update.effective_chat.id,
                                     action=telegram.ChatAction.UPLOAD_DOCUMENT)
        context.bot.send_document(chat_id=update.effective_chat.id, document=f)
    os.remove(path_nome_file)
    return ConversationHandler.END


def annulla(update, context):
    context.bot.edit_message_text(text="❌ *Annullato con successo.*",
                                  chat_id=update.effective_chat.id,
                                  message_id=update.callback_query.message.message_id,
                                  parse_mode=telegram.ParseMode.MARKDOWN)
    return ConversationHandler.END

def test(update, context):
    with open("test.mp4", "rb") as f:
        context.bot.send_document(chat_id=update.effective_chat.id, document=f)


dispatcher.add_handler(CommandHandler("privacy", privacy))
dispatcher.add_handler(CommandHandler("test", test))
dispatcher.add_handler(CommandHandler("set", api_id))
dispatcher.add_handler(
    ConversationHandler(
        entry_points=[CommandHandler("start", ask_for_category),
                      CallbackQueryHandler(ask_for_category, pattern="^2")],
        states={
            ASK_TARGET: [CallbackQueryHandler(ask_for_target, pattern="^1")],
            ASK_FILE: [CallbackQueryHandler(ask_for_file, pattern="[a-z]")],
            CONVERT: [MessageHandler(Filters.video | Filters.photo | Filters.document | Filters.audio,
                                     convert)]
        },
        fallbacks=[CommandHandler("start", ask_for_category),
                   CallbackQueryHandler(ask_for_target, pattern="^1")
                   ],
        allow_reentry=True
    )
)

dispatcher.add_handler(ConversationHandler(
    entry_points=[CommandHandler('new', ask_pdfdocx), CallbackQueryHandler(ask_pdfdocx, pattern="^0")],
    states={
        ASK_FILENAME: [CallbackQueryHandler(ask_filename, pattern="^4")],
        ASK_TITLE: [MessageHandler(Filters.text & ~Filters.command, ask_title)],
        ASK_CONTENT: [MessageHandler((Filters.text | Filters.photo) & ~Filters.command, ask_content)]
    },
    fallbacks=[CommandHandler('done', done), CallbackQueryHandler(annulla, pattern="3")],
    allow_reentry=True
))

updater.start_polling()
updater.idle()
